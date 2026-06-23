from __future__ import annotations

from datetime import date
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "docs"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

TODAY = date(2026, 6, 13)
DATE_LABEL = TODAY.strftime("%B %-d, %Y")

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
INK = RGBColor(0x1E, 0x2A, 0x35)
MUTED = RGBColor(0x5F, 0x6B, 0x76)
LIGHT_FILL = "F2F4F7"
PALE_BLUE = "E8EEF5"
WHITE = "FFFFFF"
BORDER = "C9D1D9"


def set_run_font(
    run,
    *,
    name: str = "Calibri",
    size: float | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = BORDER, size: int = 6) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa: Sequence[int], indent_dxa: int = 120) -> None:
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, width in enumerate(widths_dxa):
            cell = row.cells[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def style_block_paragraph(paragraph, *, fill: str, border_color: str, border_size: int = 6) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = p_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        p_pr.append(shading)
    shading.set(qn("w:fill"), fill)

    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        border = borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), str(border_size))
        border.set(qn("w:space"), "5")
        border.set(qn("w:color"), border_color)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])


def configure_document(doc: Document, running_label: str) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in (
        ("Title", 24, INK, 0, 4),
        ("Subtitle", 13, MUTED, 0, 14),
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = style_name.startswith("Heading") or style_name == "Title"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_p.paragraph_format.space_after = Pt(0)
    run = header_p.add_run(running_label)
    set_run_font(run, size=9, color=MUTED, bold=True)

    footer = section.footer
    footer_p = footer.paragraphs[0]
    add_page_number(footer_p)


def add_title_block(
    doc: Document,
    title: str,
    subtitle: str,
    doc_type: str,
    status: str = "Planning baseline",
) -> None:
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(8)
    kicker.paragraph_format.space_after = Pt(4)
    run = kicker.add_run(doc_type.upper())
    set_run_font(run, size=9, color=BLUE, bold=True)

    title_p = doc.add_paragraph(style="Title")
    title_p.add_run(title)

    subtitle_p = doc.add_paragraph(style="Subtitle")
    subtitle_p.add_run(subtitle)

    rows = [
        ("Project", "Chess App (working title)"),
        ("Version", "1.0"),
        ("Date", DATE_LABEL),
        ("Status", status),
    ]
    for label, value in rows:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        label_run = p.add_run(f"{label}: ")
        set_run_font(label_run, size=10.5, bold=True, color=INK)
        value_run = p.add_run(value)
        set_run_font(value_run, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_body(doc: Document, text: str, *, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        first, rest = text[: len(bold_lead)], text[len(bold_lead) :]
        run = p.add_run(first)
        set_run_font(run, bold=True)
        p.add_run(rest)
    else:
        p.add_run(text)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_callout(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    style_block_paragraph(p, fill=PALE_BLUE, border_color="B8C8D8")
    label_run = p.add_run(f"{label}: ")
    set_run_font(label_run, bold=True, color=DARK_BLUE)
    p.add_run(text)


def add_table(
    doc: Document,
    headers: Sequence[str],
    rows: Sequence[Sequence[str]],
    widths_dxa: Sequence[int],
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_FILL)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        set_run_font(run, size=10, bold=True, color=INK)
    for row_data in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_data):
            cell = cells[index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            set_run_font(run, size=9.5, color=INK)
    set_table_geometry(table, widths_dxa)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_code_block(doc: Document, lines: Sequence[str]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.0
    style_block_paragraph(p, fill="F6F8FA", border_color="D8DEE4", border_size=4)
    for index, line in enumerate(lines):
        run = p.add_run(line)
        set_run_font(run, name="Menlo", size=8.5, color=INK)
        if index < len(lines) - 1:
            run.add_break()


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


def build_prd() -> Path:
    doc = Document()
    configure_document(doc, "Chess App | Product Requirements")
    add_title_block(
        doc,
        "Product Requirements Document",
        "Authoritative chess gameplay, online multiplayer, and adaptive AI coaching",
        "Product requirements",
    )

    add_callout(
        doc,
        "Product direction",
        "Build a native iOS chess trainer that is immediately playable, becomes social through reliable online games, and differentiates through an adaptive AI sparring partner with concise post-game coaching.",
    )

    doc.add_heading("1. Executive Summary", level=1)
    add_body(
        doc,
        "The product is a native SwiftUI chess application backed by a privately operated service deployed through Coolify. "
        "The backend is authoritative for rules, game state, clocks, results, and persistence. The initial release prioritizes "
        "single-player play and friend matches before ranked matchmaking or an automated machine-learning improvement loop."
    )
    add_body(
        doc,
        "The central product promise is not merely online chess. The intended differentiator is a practical training experience: "
        "an AI opponent that adapts to the player, identifies recurring weaknesses, and gives short explanations after games."
    )

    doc.add_heading("2. Problem and Opportunity", level=1)
    add_body(
        doc,
        "Chess players can already choose from mature playing platforms. A new product therefore needs a focused reason to exist. "
        "Many beginner and improving players want useful practice without long analysis sessions, dense engine notation, or the pressure "
        "of public ranked play. The opportunity is to combine a polished native game experience with personalized, understandable coaching."
    )
    add_bullets(
        doc,
        [
            "Beginners need legal, responsive gameplay and feedback they can understand.",
            "Improving players need opponents that target their weaknesses instead of only maximizing playing strength.",
            "Casual players need low-friction friend games and reliable reconnect behavior.",
            "The product team needs complete, structured game records to support review, analytics, and future model development.",
        ],
    )

    doc.add_heading("3. Product Vision and Principles", level=1)
    add_body(
        doc,
        "Vision: A chess trainer that plays against you, learns your weaknesses, and gives short practical explanations after each game."
    )
    add_table(
        doc,
        ["Principle", "Implication"],
        [
            ("Server authority", "Clients submit commands; the backend validates and publishes canonical state."),
            ("Playable before intelligent", "Random, heuristic, and shallow-search bots ship before a custom ML model."),
            ("Training data is earned", "Raw user games are retained, filtered, labeled, and evaluated before use in training."),
            ("Reliable core first", "Reconnects, idempotency, clocks, and persistence precede ranked systems and social features."),
            ("Native product quality", "SwiftUI remains the primary client; Unity is excluded unless future visuals require it."),
        ],
        [2300, 7060],
    )

    doc.add_heading("4. Target Users", level=1)
    doc.add_heading("4.1 Primary Persona: Improving Casual Player", level=2)
    add_bullets(
        doc,
        [
            "Knows the rules but misses tactics and strategic plans.",
            "Prefers short games and concise guidance over deep engine analysis.",
            "Wants difficulty that remains challenging without feeling punitive.",
        ],
    )
    doc.add_heading("4.2 Secondary Persona: Friend-Game Player", level=2)
    add_bullets(
        doc,
        [
            "Wants to create a private match and share a simple code or link.",
            "Expects moves and clocks to remain synchronized across interruptions.",
            "Values game history and replay more than competitive ranking.",
        ],
    )
    doc.add_heading("4.3 Internal Persona: Product and ML Operator", level=2)
    add_bullets(
        doc,
        [
            "Needs searchable games, moves, outcomes, model versions, and analysis status.",
            "Needs controlled model promotion based on tests against fixed baselines.",
            "Needs auditability for game outcomes, suspicious behavior, and data provenance.",
        ],
    )

    doc.add_heading("5. Goals and Non-Goals", level=1)
    add_table(
        doc,
        ["Goals for Initial Product", "Explicitly Deferred"],
        [
            ("Complete legal chess games against server-backed AI.", "Clubs, tournaments, teams, and social feeds."),
            ("Private friend games with real-time moves and reconnects.", "Open chat, direct messaging, and moderation systems."),
            ("Persistent game history with PGN/FEN replay data.", "Advanced puzzle generation and a full course marketplace."),
            ("Basic accounts, player profile, and operational telemetry.", "Subscriptions, advertising, and complex entitlement tiers."),
            ("A safe path to future matchmaking and ratings.", "Unfiltered training on all user games or automatic model promotion."),
        ],
        [4680, 4680],
    )

    add_page_break(doc)
    doc.add_heading("6. MVP Scope", level=1)
    doc.add_heading("6.1 MVP 1: Single-Player Foundation", level=2)
    add_bullets(
        doc,
        [
            "Native chess board with selection, legal-move highlighting, promotion, check, checkmate, stalemate, draw, resign, and restart.",
            "Backend game creation and authoritative move validation using python-chess.",
            "Server-backed AI with at least random and simple heuristic difficulty.",
            "Game persistence, PGN generation, current FEN, move history, and local-friendly error states.",
            "Basic game history and replay.",
        ],
    )
    doc.add_heading("6.2 MVP 2: Online Friend Match", level=2)
    add_bullets(
        doc,
        [
            "Create a private game and receive an invite code.",
            "Join by code, assign colors, and begin when both players are ready.",
            "Receive real-time authoritative state over WebSocket.",
            "Reconnect to an active game and recover current state.",
            "Resign and record final result; basic clock support may be introduced here or in MVP 3.",
        ],
    )
    doc.add_heading("6.3 Post-MVP Sequence", level=2)
    add_numbered(
        doc,
        [
            "Quick matchmaking with basic rating and time controls.",
            "Post-game review and engine-labeled coaching.",
            "Controlled model-data preparation and offline evaluation.",
            "Adaptive difficulty and personalized weakness tracking.",
        ],
    )

    doc.add_heading("7. Core User Journeys", level=1)
    doc.add_heading("7.1 Play Against AI", level=2)
    add_numbered(
        doc,
        [
            "Player chooses difficulty and color.",
            "App creates an AI game through HTTPS and receives canonical initial state.",
            "Player makes a move locally; the app submits the UCI command.",
            "Backend validates the move, updates state, persists it, and requests an AI move when appropriate.",
            "Backend validates the AI response before applying it.",
            "App renders the authoritative state until the game ends.",
            "Player can review the completed move list and result.",
        ],
    )
    doc.add_heading("7.2 Play a Friend", level=2)
    add_numbered(
        doc,
        [
            "Host creates a private game and shares an invite code.",
            "Guest joins; both clients establish authenticated WebSocket sessions.",
            "Server assigns colors and broadcasts initial state.",
            "Each move is submitted as a command and accepted or rejected by the server.",
            "On interruption, the client reconnects and requests a state snapshot from the last known sequence number.",
            "The server finalizes result and stores PGN after checkmate, draw, resignation, or timeout.",
        ],
    )

    doc.add_heading("8. Functional Requirements", level=1)
    add_table(
        doc,
        ["ID", "Requirement", "Priority"],
        [
            ("FR-01", "Create a chess game with variant, participants, color assignment, and optional time control.", "Must"),
            ("FR-02", "Validate every submitted move against canonical server state.", "Must"),
            ("FR-03", "Return canonical FEN, move, turn, status, result, and monotonic sequence number.", "Must"),
            ("FR-04", "Reject illegal, stale, duplicate, unauthorized, and out-of-turn commands without corrupting state.", "Must"),
            ("FR-05", "Persist moves in UCI and SAN, FEN after move, timestamps, and actor.", "Must"),
            ("FR-06", "Generate and retain a complete PGN for finished games.", "Must"),
            ("FR-07", "Support an AI move request through an internal service contract and validate its response.", "Must"),
            ("FR-08", "Create and join private games using expiring invite codes.", "Must for MVP 2"),
            ("FR-09", "Broadcast accepted state changes to both players over WebSocket.", "Must for MVP 2"),
            ("FR-10", "Recover active state after reconnect or app relaunch.", "Must for MVP 2"),
            ("FR-11", "Expose game history and replay data for the current user.", "Should"),
            ("FR-12", "Record model version and difficulty for every AI game.", "Should"),
        ],
        [900, 7110, 1350],
    )

    doc.add_heading("9. Authoritative Protocol", level=1)
    add_body(
        doc,
        "The client sends commands, not state mutations. Every command includes a client-generated command ID and the sequence number "
        "the client believes it is extending. The server emits a new sequence number only after the transition is committed."
    )
    add_code_block(
        doc,
        [
            "{",
            '  "type": "move",',
            '  "command_id": "018f...uuid",',
            '  "game_id": "abc123",',
            '  "expected_sequence": 12,',
            '  "move": "e2e4"',
            "}",
        ],
    )
    add_code_block(
        doc,
        [
            "{",
            '  "type": "state_updated",',
            '  "game_id": "abc123",',
            '  "sequence": 13,',
            '  "fen": "...",',
            '  "last_move": {"uci": "e2e4", "san": "e4"},',
            '  "turn": "black",',
            '  "status": "active",',
            '  "result": null',
            "}",
        ],
    )
    add_callout(
        doc,
        "Protocol decision",
        "Legal moves may be returned on demand or for selected pieces. Sending the full legal-move list after every move is optional and should be measured because it increases payload size and couples UI hints to network latency.",
    )

    doc.add_heading("10. System Architecture", level=1)
    add_code_block(
        doc,
        [
            "SwiftUI iOS App",
            "  | HTTPS + WebSocket",
            "  v",
            "FastAPI Application",
            "  |-- Auth and player profile",
            "  |-- Game command service (python-chess)",
            "  |-- Matchmaking and invite service",
            "  |-- WebSocket connection manager",
            "  |-- AI service client",
            "  |-- Background job producer",
            "  |",
            "  |-- PostgreSQL: durable users, games, moves, ratings",
            "  `-- Redis: queues, presence, locks, transient active state",
            "",
            "Worker / AI Services",
            "  |-- analysis and training-sample preparation",
            "  `-- model inference and optional private engine analysis",
        ],
    )
    add_body(
        doc,
        "For the first scaffold, the backend should remain a modular monolith. Auth, game, invite, and matchmaking boundaries should be "
        "clear in code, but separate deployable microservices are not required until scaling or ownership justifies them."
    )

    doc.add_heading("11. Data Model", level=1)
    add_table(
        doc,
        ["Entity", "Minimum Fields"],
        [
            ("users", "id, apple_user_id, username, rating, created_at, updated_at"),
            ("games", "id, white_user_id, black_user_id, mode, status, result, current_fen, pgn, time_control, sequence, timestamps"),
            ("moves", "id, game_id, ply, player_id, uci_move, san_move, fen_after, command_id, created_at"),
            ("ratings", "id, user_id, game_id, rating_before, rating_after, delta, system, created_at"),
            ("ai_games", "id, game_id, user_id, model_version, difficulty, result, created_at"),
            ("training_samples", "id, source_game_id, ply, fen, move_played, engine_score, label_version, split, created_at"),
        ],
        [2100, 7260],
    )
    add_body(
        doc,
        "Schema refinements expected during implementation include unique command IDs per game, immutable move ordering, optimistic or "
        "pessimistic concurrency controls, and separate rating-event records instead of storing only a mutable rating on the user."
    )

    doc.add_heading("12. Non-Functional Requirements", level=1)
    add_table(
        doc,
        ["Area", "Initial Requirement"],
        [
            ("Correctness", "No accepted command may create an illegal board state; move order and result must be reproducible from stored moves."),
            ("Latency", "Target p95 accepted-move response below 300 ms excluding AI think time under expected MVP load."),
            ("Availability", "Graceful client recovery after transient disconnect; no active game should depend on one process's memory."),
            ("Consistency", "Commands are idempotent and sequenced; stale commands return canonical recovery state."),
            ("Security", "Authenticated authorization on every game command; secrets remain server-side; rate limits protect creation and sockets."),
            ("Privacy", "Collect only required account and gameplay data; define retention and deletion behavior before public release."),
            ("Observability", "Structured logs, request IDs, game IDs, command IDs, metrics, error tracking, and health checks."),
            ("Portability", "Backend contracts support future Android and web clients without iOS-specific semantics."),
        ],
        [1900, 7460],
    )

    doc.add_heading("13. AI and Training Data Requirements", level=1)
    add_bullets(
        doc,
        [
            "The app never communicates directly with a model or chess engine.",
            "The game service sends FEN, difficulty, constraints, and correlation metadata to the AI service.",
            "The game service validates every returned move with python-chess.",
            "Model version, latency, selected move, and failure mode are recorded.",
            "Raw games are not automatically converted into training samples.",
            "Training samples require provenance, filtering, labeling, evaluation, and deterministic train/test splits.",
            "A candidate model is deployed only after it passes fixed correctness and playing-strength baselines.",
        ],
    )
    add_callout(
        doc,
        "Licensing",
        "Stockfish is distributed under GPL v3. Server-side private use can simplify distribution concerns, but deployment and source-offer obligations must be reviewed before Stockfish becomes part of a commercial production service or distributed client.",
    )

    doc.add_heading("14. Anti-Cheat and Abuse Controls", level=1)
    add_bullets(
        doc,
        [
            "The server validates all moves, clocks, game completion, and participant authorization.",
            "Game creation, invite attempts, API calls, and WebSocket connections are rate-limited.",
            "Move timing and reconnect events are retained for later anomaly analysis.",
            "No chat is included in the MVP, avoiding an immediate moderation surface.",
            "Accuracy-based flags are advisory and reviewed; they do not automatically punish users in the initial product.",
        ],
    )

    doc.add_heading("15. Success Measures", level=1)
    add_table(
        doc,
        ["Stage", "Primary Measures"],
        [
            ("Internal alpha", "Crash-free games, legal-state correctness, AI response success, reconnect recovery, test pass rate."),
            ("MVP 1", "Game start-to-finish rate, average games per active user, resign/restart rate, backend p95 latency."),
            ("MVP 2", "Invite completion rate, disconnect rate, successful reconnect rate, friend games completed."),
            ("Coaching release", "Review completion rate, repeat play after feedback, explanation helpfulness, skill retention proxies."),
        ],
        [2100, 7260],
    )

    doc.add_heading("16. Release Acceptance Criteria", level=1)
    add_bullets(
        doc,
        [
            "All standard legal moves, castling, en passant, promotion, checkmate, stalemate, repetition, fifty-move rule, and insufficient material are covered by automated tests.",
            "Duplicate and stale commands cannot create duplicate moves or regress sequence state.",
            "A completed game can be reconstructed from persisted records and exports valid PGN.",
            "AI failures time out cleanly and do not leave the game locked or corrupted.",
            "An interrupted friend game resumes from canonical state without manual database intervention.",
            "Coolify health checks, migrations, backups, and rollback instructions are documented and exercised in staging.",
            "No critical or high-severity security findings remain open for the release scope.",
        ],
    )

    doc.add_heading("17. Risks and Mitigations", level=1)
    add_table(
        doc,
        ["Risk", "Mitigation"],
        [
            ("Multiplayer state races", "Sequence commands, enforce idempotency, serialize per-game transitions, and test concurrent submissions."),
            ("Redis treated as durable truth", "Keep durable game and move records in PostgreSQL; use Redis for coordination and recoverable transient state."),
            ("Weak custom model delays product", "Ship deterministic baseline bots and maintain a stable AI service interface."),
            ("Scope expansion", "Use phase gates; defer ranked, chat, tournaments, subscriptions, and advanced ML."),
            ("Engine/data licensing", "Track software and dataset licenses; obtain legal review before distribution or commercial deployment."),
            ("Operational complexity", "Begin with a modular monolith plus worker, PostgreSQL, and Redis; split services only on measured need."),
        ],
        [2800, 6560],
    )

    doc.add_heading("18. Open Decisions", level=1)
    add_bullets(
        doc,
        [
            "Product name, visual identity, and exact minimum iOS version.",
            "Authentication launch method: Sign in with Apple only, guest accounts, or both.",
            "Whether clocks launch in friend matches or with matchmaking.",
            "Rating system choice and whether provisional ratings are visible.",
            "Initial hosting topology and expected concurrent-game capacity.",
            "Data retention, account deletion, privacy policy, and regional requirements.",
            "Commercial use and deployment posture for Stockfish or other GPL components.",
        ],
    )

    path = OUTPUT_DIR / "Chess_App_Product_Requirements.docx"
    doc.save(path)
    return path


def build_phase_plan() -> Path:
    doc = Document()
    configure_document(doc, "Chess App | Phased Build Plan")
    add_title_block(
        doc,
        "Phased Build Plan",
        "Implementation sequence from empty repository to adaptive chess trainer",
        "Engineering delivery plan",
    )
    add_callout(
        doc,
        "Planning assumption",
        "The repository is currently empty. Estimates are directional engineering effort for one experienced full-time engineer and exclude App Store review, legal review, branding, and external model-training time.",
    )

    doc.add_heading("1. Delivery Strategy", level=1)
    add_body(
        doc,
        "The build proceeds through vertical slices. Each phase must leave the product runnable, tested, and deployable. The backend begins "
        "as a modular FastAPI application with a separate worker and AI service contract. PostgreSQL owns durable state; Redis supports "
        "coordination, presence, queues, and locks. The iOS app adopts server authority early so local-only assumptions do not leak into multiplayer."
    )
    add_bullets(
        doc,
        [
            "Correctness before scale: prove chess state transitions and persistence first.",
            "Contracts before clients: stabilize API and WebSocket message envelopes with tests.",
            "Single process before distributed complexity: introduce Redis-backed coordination only where a phase needs it.",
            "Baseline AI before ML: preserve the product schedule even when a custom model is not ready.",
            "Phase gates are mandatory: deferred features do not enter a phase without replacing an existing commitment.",
        ],
    )

    doc.add_heading("2. Target Repository Structure", level=1)
    add_code_block(
        doc,
        [
            "chess-app/",
            "  apps/",
            "    ios/                       # SwiftUI application and Xcode project",
            "  services/",
            "    api/                       # FastAPI modular monolith",
            "      app/",
            "        api/                   # HTTP and WebSocket transport",
            "        auth/                  # identity and authorization",
            "        games/                 # domain model and commands",
            "        matchmaking/           # invites and queues",
            "        ai/                    # internal AI client/contract",
            "        db/                    # sessions, models, migrations",
            "        core/                  # settings, logging, telemetry",
            "      tests/",
            "    worker/                    # background analysis jobs",
            "    ai/                        # inference service and baseline bots",
            "  deploy/",
            "    compose.yaml               # local and Coolify service definitions",
            "    coolify/                   # environment and runbook notes",
            "  docs/",
            "  .github/workflows/",
        ],
    )

    doc.add_heading("3. Technology Baseline", level=1)
    add_table(
        doc,
        ["Layer", "Initial Choice", "Reason"],
        [
            ("iOS", "Swift 6, SwiftUI, async/await", "Native UX, App Store integration, clear state and concurrency model."),
            ("API", "Python 3.13, FastAPI, Pydantic", "Strong python-chess fit and native async WebSocket support."),
            ("Persistence", "PostgreSQL, SQLAlchemy 2, Alembic", "Transactional durable state and explicit migrations."),
            ("Coordination", "Redis", "Matchmaking queues, presence, locks, pub/sub, and transient state."),
            ("Rules", "python-chess", "Mature legal move, outcome, FEN, SAN, and PGN support."),
            ("Testing", "pytest, async test client, Testcontainers where practical", "Fast domain tests plus realistic integration coverage."),
            ("Delivery", "Docker, Compose, Coolify, GitHub Actions", "Reproducible local/staging/production deployment."),
        ],
        [1500, 2700, 5160],
    )
    add_body(
        doc,
        "Versions should be pinned during scaffolding after confirming current supported releases. The plan intentionally names technologies, not unverified version numbers."
    )

    doc.add_heading("4. Phase Overview", level=1)
    add_table(
        doc,
        ["Phase", "Outcome", "Indicative Effort"],
        [
            ("0. Foundation", "Repository, CI, containers, configuration, observability baseline.", "3-5 days"),
            ("1. Chess Core", "Authoritative game creation, legal moves, persistence, PGN.", "1-2 weeks"),
            ("2. Native Board", "Playable SwiftUI board connected to backend game APIs.", "2-3 weeks"),
            ("3. AI Opponent", "Server-backed baseline bots with difficulty and resilience.", "1-2 weeks"),
            ("4. Friend Multiplayer", "Invite games, WebSockets, reconnects, presence.", "2-4 weeks"),
            ("5. Matchmaking", "Queues, clocks, ratings, records, abuse limits.", "2-4 weeks"),
            ("6. Review", "Game history, replay, analysis jobs, concise feedback.", "2-3 weeks"),
            ("7. ML Pipeline", "Filtered datasets, evaluation harness, model registry.", "3-6 weeks"),
            ("8. Adaptive Coach", "Personalized difficulty and weakness-oriented coaching.", "3-6+ weeks"),
        ],
        [1700, 5860, 1800],
    )

    add_page_break(doc)
    doc.add_heading("5. Phase 0 - Foundation and Scaffold", level=1)
    add_body(doc, "Objective: Turn the empty workspace into a reproducible development and deployment baseline.")
    doc.add_heading("Deliverables", level=2)
    add_bullets(
        doc,
        [
            "Git repository with documented branch and pull-request conventions.",
            "FastAPI service package with typed settings, health/readiness endpoints, structured logging, and request IDs.",
            "PostgreSQL and Redis containers with local Compose configuration and isolated test configuration.",
            "SQLAlchemy session management and an initial Alembic migration.",
            "pytest setup, linting, formatting, type checking, and a minimal GitHub Actions workflow.",
            "Production Dockerfile, non-root runtime user, graceful shutdown, and Coolify environment template.",
            "Architecture decision records for server authority, PostgreSQL durability, Redis responsibilities, and modular-monolith boundaries.",
        ],
    )
    doc.add_heading("Verification", level=2)
    add_bullets(
        doc,
        [
            "One command starts API, PostgreSQL, and Redis locally.",
            "Health check verifies process liveness; readiness verifies database and Redis connectivity.",
            "Migration upgrade and downgrade run in CI against a clean database.",
            "Static checks and tests pass without local-only dependencies.",
            "Container image builds and starts with production settings.",
        ],
    )
    add_callout(
        doc,
        "Phase 0 exit gate",
        "A clean checkout can be configured, tested, migrated, and started from documented commands. No gameplay behavior is required yet.",
    )

    doc.add_heading("6. Phase 1 - Authoritative Chess Core", level=1)
    add_body(doc, "Objective: Complete a legal chess game through backend APIs with durable, reconstructable state.")
    doc.add_heading("Deliverables", level=2)
    add_bullets(
        doc,
        [
            "Game aggregate with participants, mode, status, result, FEN, sequence, and timestamps.",
            "Create-game, get-game, submit-move, resign, and list-moves endpoints.",
            "Move command contract with command ID and expected sequence.",
            "python-chess adapter for UCI validation, SAN generation, outcomes, FEN, and PGN.",
            "Transactional move append with unique idempotency constraint and per-game concurrency strategy.",
            "Database models and migrations for users, games, and moves; development identity fixture.",
            "Domain tests for standard moves and every special outcome.",
        ],
    )
    doc.add_heading("Verification", level=2)
    add_bullets(
        doc,
        [
            "Automated API test plays a complete game to checkmate and validates final PGN.",
            "Illegal, out-of-turn, stale, duplicate, and unauthorized commands are rejected deterministically.",
            "Two simultaneous move submissions cannot both advance the same sequence.",
            "Game state can be rebuilt from persisted moves and matches stored FEN.",
            "No Redis dependency is required for correctness in this phase.",
        ],
    )
    add_callout(
        doc,
        "Phase 1 exit gate",
        "A test client can create and finish valid games entirely through the API; persisted records reproduce every final position.",
    )

    doc.add_heading("7. Phase 2 - Native iOS Board and Game Shell", level=1)
    add_body(doc, "Objective: Deliver a polished native board that consumes authoritative backend state.")
    doc.add_heading("Deliverables", level=2)
    add_bullets(
        doc,
        [
            "SwiftUI app shell with dependency injection, environment configuration, and feature-based organization.",
            "Chess board rendering, orientation, piece assets, selection, move highlighting, and promotion UI.",
            "Game state model that distinguishes local interaction state from authoritative server state.",
            "Typed API client using async/await, cancellation, retry boundaries, and explicit errors.",
            "Game screen controls for resign, restart/new game, loading, offline, and terminal outcomes.",
            "Accessibility labels, dynamic layout checks, and primary state previews.",
        ],
    )
    doc.add_heading("Verification", level=2)
    add_bullets(
        doc,
        [
            "The board renders correctly on supported phone sizes and both orientations if enabled.",
            "Legal selections and promotions are usable with VoiceOver labels.",
            "The app does not commit an unconfirmed move as canonical state.",
            "Network failures preserve a recoverable UI and allow state refresh.",
            "A user can complete a two-sided test game through the app against the Phase 1 API.",
        ],
    )
    add_callout(
        doc,
        "Phase 2 exit gate",
        "The native client can reliably play and display a backend-owned game from start to completion.",
    )

    doc.add_heading("8. Phase 3 - Single-Player AI", level=1)
    add_body(doc, "Objective: Make the product independently playable without waiting for a custom ML model.")
    doc.add_heading("Deliverables", level=2)
    add_bullets(
        doc,
        [
            "Internal AI request/response contract using FEN, difficulty, constraints, and correlation ID.",
            "Baseline bot levels: random legal move, heuristic capture/check preference, and shallow search.",
            "AI game orchestration that requests a move only after the human transition commits.",
            "Validation of every AI-returned move before application.",
            "Timeout, retry, fallback, and circuit-breaker behavior that cannot corrupt a game.",
            "Model/bot version and think-time telemetry stored with each AI game.",
        ],
    )
    doc.add_heading("Verification", level=2)
    add_bullets(
        doc,
        [
            "Each difficulty returns legal moves across a fixed position suite.",
            "Invalid and timed-out AI responses produce controlled recovery or fallback.",
            "Concurrent retries do not apply multiple AI moves.",
            "A user can complete multiple AI games and see them in history.",
        ],
    )
    add_callout(
        doc,
        "Phase 3 exit gate",
        "MVP 1 is release-candidate quality: users can play, resign, restart, and review games against at least two distinct baseline difficulty levels.",
    )

    doc.add_heading("9. Phase 4 - Friend-Code Multiplayer", level=1)
    add_body(doc, "Objective: Support reliable private real-time games between two authenticated clients.")
    doc.add_heading("Deliverables", level=2)
    add_bullets(
        doc,
        [
            "Sign in with Apple or a documented temporary identity path for pre-release testing.",
            "Private game creation, short expiring invite codes, join validation, and color assignment.",
            "Authenticated WebSocket endpoint with connection lifecycle, heartbeat, and message envelopes.",
            "Redis-backed presence and cross-instance pub/sub; PostgreSQL remains the durable game source.",
            "Snapshot plus sequence-based recovery after reconnect or missed events.",
            "Client WebSocket actor/service with backoff, cancellation, foreground/background handling, and resubscription.",
        ],
    )
    doc.add_heading("Verification", level=2)
    add_bullets(
        doc,
        [
            "Two devices or simulators complete a private game with synchronized state.",
            "Disconnecting either client for a defined interval does not lose committed moves.",
            "A client missing events recovers from a canonical snapshot and current sequence.",
            "An unauthorized user cannot observe or command a private game.",
            "Multi-instance test proves events cross API process boundaries through Redis.",
        ],
    )
    add_callout(
        doc,
        "Phase 4 exit gate",
        "MVP 2 is release-candidate quality: invite creation, joining, real-time play, and reconnect recovery work across multiple backend instances.",
    )

    doc.add_heading("10. Phase 5 - Matchmaking, Clocks, and Ratings", level=1)
    add_body(doc, "Objective: Add fair quick play without weakening the correctness established in friend games.")
    doc.add_heading("Deliverables", level=2)
    add_bullets(
        doc,
        [
            "Redis matchmaking queues partitioned by time control and rating band.",
            "Atomic pairing, cancellation, duplicate-entry prevention, and abandoned-match cleanup.",
            "Server-owned clocks with latency policy, timeout adjudication, and persisted timing events.",
            "Initial rating system with immutable rating events and provisional handling.",
            "Win/loss/draw records, basic profile display, and rate limits.",
            "Operational controls for queue depth, suspicious usage, and incident recovery.",
        ],
    )
    doc.add_heading("Verification", level=2)
    add_bullets(
        doc,
        [
            "Load test pairs expected concurrent users without duplicate matches.",
            "Clock outcomes remain correct through disconnects and process restarts.",
            "Rating updates are transactional with game finalization and cannot apply twice.",
            "Queue cancellation and timeout leave no stuck users or orphaned games.",
        ],
    )
    add_callout(
        doc,
        "Phase 5 exit gate",
        "Quick play can be monitored and operated safely at the agreed launch concurrency, with deterministic game and rating outcomes.",
    )

    doc.add_heading("11. Phase 6 - History, Review, and Coaching Baseline", level=1)
    add_body(doc, "Objective: Turn completed games into understandable learning moments.")
    doc.add_heading("Deliverables", level=2)
    add_bullets(
        doc,
        [
            "Paginated game history, replay controls, PGN export, and position navigation.",
            "Background analysis jobs with explicit status, retry, and dead-letter behavior.",
            "Engine-evaluation adapter isolated from the public game request path.",
            "Detection of major swings, missed tactics, and key moments.",
            "Short template-driven explanations with clear uncertainty and no fabricated claims.",
            "Internal review tools for analysis quality and failed jobs.",
        ],
    )
    doc.add_heading("Verification", level=2)
    add_bullets(
        doc,
        [
            "Analysis never blocks game completion or corrupts game records.",
            "Jobs are idempotent and can be replayed after worker failure.",
            "Feedback references the correct position and legal alternatives.",
            "The app displays analysis progress, absence, and failure honestly.",
        ],
    )
    add_callout(
        doc,
        "Phase 6 exit gate",
        "A completed game reliably produces a concise review that a beginner can connect to specific replay positions.",
    )

    doc.add_heading("12. Phase 7 - Controlled ML Improvement Pipeline", level=1)
    add_body(doc, "Objective: Produce trustworthy training candidates without allowing raw gameplay to degrade the model.")
    doc.add_heading("Deliverables", level=2)
    add_bullets(
        doc,
        [
            "Immutable raw-game export with provenance, consent/retention controls, and schema version.",
            "Filtering rules for corruption, low-signal positions, duplicates, and policy exclusions.",
            "Reference-engine labeling jobs and reproducible feature generation.",
            "Dataset versioning, train/validation/test split protection, and leakage checks.",
            "Evaluation harness covering legality, tactical suites, baseline match play, latency, and resource use.",
            "Model registry with candidate, approved, deployed, and retired states.",
            "Manual promotion and rollback; no automatic production promotion.",
        ],
    )
    doc.add_heading("Verification", level=2)
    add_bullets(
        doc,
        [
            "A dataset can be regenerated from source records and versioned configuration.",
            "Evaluation produces repeatable results against fixed seeds and baselines.",
            "A candidate that regresses legality, strength, latency, or stability cannot be promoted.",
            "Production inference can be rolled back to a prior model without a client update.",
        ],
    )
    add_callout(
        doc,
        "Phase 7 exit gate",
        "One candidate model completes the full offline pipeline and is promoted only after documented comparison with baseline bots.",
    )

    doc.add_heading("13. Phase 8 - Adaptive AI Coach", level=1)
    add_body(doc, "Objective: Deliver the product's differentiator through personalized practice and feedback.")
    doc.add_heading("Deliverables", level=2)
    add_bullets(
        doc,
        [
            "Player weakness profile derived from reviewed games and bounded behavioral signals.",
            "Adaptive opponent selection or move policy that targets weaknesses while maintaining a fair experience.",
            "Difficulty calibration with measurable target win-rate or challenge bands.",
            "Personalized post-game summaries and suggested practice themes.",
            "Experiment framework with privacy-conscious assignment and guardrails.",
            "User controls for personalization and data deletion.",
        ],
    )
    doc.add_heading("Verification", level=2)
    add_bullets(
        doc,
        [
            "Personalization improves agreed learning or engagement measures without increasing illegal or unstable behavior.",
            "Difficulty changes are explainable, bounded, and reversible.",
            "Users with insufficient data receive a stable non-personalized experience.",
            "Deleting an account removes or irreversibly anonymizes associated personalization data according to policy.",
        ],
    )
    add_callout(
        doc,
        "Phase 8 exit gate",
        "The adaptive experience demonstrates a measurable benefit over static difficulty in a controlled test and meets privacy requirements.",
    )

    doc.add_heading("14. Cross-Phase Quality Plan", level=1)
    add_table(
        doc,
        ["Test Layer", "Required Coverage"],
        [
            ("Domain", "Move legality, outcomes, clocks, ratings, command sequencing, idempotency, and state reconstruction."),
            ("API contract", "Success and error envelopes, auth, pagination, versioning, stale-state recovery, and schema compatibility."),
            ("Integration", "PostgreSQL transactions, migrations, Redis locks/pub-sub/queues, worker retries, and AI failures."),
            ("iOS", "State reducers/models, API decoding, WebSocket recovery, board interactions, accessibility, and snapshot/UI smoke tests."),
            ("End to end", "AI game, friend game, reconnect, timeout, result persistence, history, and analysis."),
            ("Operational", "Container startup, health checks, backup restore, migration rollback, load, and deploy rollback."),
        ],
        [2000, 7360],
    )

    doc.add_heading("15. Coolify Deployment Plan", level=1)
    add_table(
        doc,
        ["Service", "Phase Introduced", "Deployment Responsibility"],
        [
            ("chess-api", "0", "HTTP/WebSocket API, migrations as controlled release step, health/readiness."),
            ("postgres", "0", "Durable records, automated backups, restore test, persistent volume."),
            ("redis", "0; required by 4", "Coordination and queues; persistence policy based on recoverability needs."),
            ("chess-ai", "3", "Baseline/custom inference behind internal network and strict timeout."),
            ("chess-worker", "6", "Analysis and data jobs with independent scaling and failure handling."),
            ("admin-dashboard", "Optional after 4", "Internal operations only; not a prerequisite for MVP."),
        ],
        [1900, 1700, 5760],
    )
    add_bullets(
        doc,
        [
            "Separate staging and production environments, secrets, databases, domains, and deployment approvals.",
            "Run migrations as an explicit, observable release action; avoid hidden startup migrations in multiple replicas.",
            "Use rolling or blue/green deployment only after WebSocket draining and compatibility behavior are tested.",
            "Backups are incomplete until a restore has been successfully exercised.",
            "Keep the prior image and migration rollback instructions available for every production release.",
        ],
    )

    doc.add_heading("16. First Scaffold Work Package", level=1)
    add_body(
        doc,
        "The immediate next implementation should cover Phase 0 plus the smallest Phase 1 vertical slice. This proves the architecture before the iOS project is generated."
    )
    add_numbered(
        doc,
        [
            "Initialize repository, backend package, tooling, Dockerfile, Compose, CI, and environment examples.",
            "Start PostgreSQL and Redis locally; add liveness/readiness checks.",
            "Create initial users, games, and moves migrations.",
            "Implement create game, retrieve game, and submit move using python-chess.",
            "Add command ID and expected sequence to the move contract.",
            "Test legal, illegal, stale, and duplicate moves plus one complete checkmate flow.",
            "Document local commands and Coolify deployment inputs.",
        ],
    )
    add_callout(
        doc,
        "Scaffold definition of done",
        "A clean environment can run the stack and execute an automated test that creates a game, submits legal moves, rejects an illegal move, and retrieves the resulting canonical state from PostgreSQL.",
    )

    doc.add_heading("17. Decision Log for Scaffolding", level=1)
    add_table(
        doc,
        ["Decision", "Baseline"],
        [
            ("Backend language", "Python with FastAPI to keep python-chess in-process for the authoritative game core."),
            ("Architecture", "Modular monolith API plus separate worker and AI deployment boundaries."),
            ("Durability", "PostgreSQL is authoritative; Redis state must be reconstructable or explicitly disposable."),
            ("Client protocol", "Command IDs, expected sequence, canonical responses, and snapshot recovery."),
            ("Identity in scaffold", "Development identity seam with production auth interface; no hard-coded auth assumptions in domain logic."),
            ("API compatibility", "Versioned `/v1` routes and explicit WebSocket message schema from the first public client."),
        ],
        [2600, 6760],
    )

    doc.add_heading("18. Deferred Until Triggered", level=1)
    add_bullets(
        doc,
        [
            "Separate auth, matchmaking, and game microservices until scale or ownership requires independent deployment.",
            "Kubernetes or bespoke orchestration while Coolify and Docker meet operational needs.",
            "GraphQL, event sourcing, and a general message bus without a demonstrated requirement.",
            "Chat, tournaments, clubs, subscriptions, and broad social features before core retention is established.",
            "Automatic anti-cheat penalties or automated model deployment without human review.",
        ],
    )

    path = OUTPUT_DIR / "Chess_App_Phased_Build_Plan.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    prd_path = build_prd()
    phase_path = build_phase_plan()
    print(prd_path)
    print(phase_path)
